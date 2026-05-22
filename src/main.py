# src/main.py

import os
import json
import io
import re
import asyncio
import httpx  # <-- NUEVO: Para consultar el RAG
import fitz  # PyMuPDF para comprimir PDFs
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Response, Depends, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse 
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF
from datetime import datetime, timedelta, timezone
from google.cloud.firestore import AsyncClient, SERVER_TIMESTAMP, Query
from google.cloud import firestore 
from google.cloud import storage 
import google.auth

# --- NUEVOS IMPORTS DEL SDK GENAI ---
from google import genai
from google.genai import types
from google.genai.errors import APIError

from src.core.security import get_current_user
from src.core.prompts import ANALYZER_SYSTEM_PROMPT, FIRST_TURN_PROMPT_TEMPLATE, FOLLOW_UP_PROMPT_TEMPLATE

# Cargar variables
load_dotenv()

# --- CONFIGURACIÓN GENAI Y STORAGE ---
try:
    raw_credentials, project_id_default = google.auth.default()
    PROJECT_ID = os.getenv("PROJECT_ID", project_id_default)
    
    from google.auth.compute_engine.credentials import Credentials as ComputeEngineCredentials
    from google.auth import impersonated_credentials
    
    if isinstance(raw_credentials, ComputeEngineCredentials):
        print("Entorno Cloud Run detectado. Impersonando cuenta para firmar URLs...")
        credentials = impersonated_credentials.Credentials(
            source_credentials=raw_credentials,
            target_principal="analize-v20@pida-ai-v20.iam.gserviceaccount.com",
            # ✅ URL LIMPIA:
            target_scopes=["https://www.googleapis.com/auth/cloud-platform"],
            lifetime=3600
        )
    else:
        credentials = raw_credentials

except Exception as e:
    print(f"Error configurando credenciales: {e}")
    PROJECT_ID = os.getenv("PROJECT_ID")
    credentials = None

LOCATION = os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1")
GEMINI_MODEL_NAME = os.getenv("GEMINI_MODEL", "gemini-2.5-pro").strip()
GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME", "pida-ai-temp-docs")

# Inicialización de Clientes
genai_client = None
storage_client = None
if PROJECT_ID:
    try:
        # Nuevo cliente de GenAI
        genai_client = genai.Client(vertexai=True, project=PROJECT_ID, location=LOCATION)
        storage_client = storage.Client(project=PROJECT_ID, credentials=credentials) 
        print(f"GenAI y Storage inicializados: {PROJECT_ID}")
    except Exception as e:
        print(f"Error inicializando GCP: {e}")

# Inicializar Firestore
db = AsyncClient(project=PROJECT_ID)

app = FastAPI(title="PIDA Document Analyzer (Streaming, GCS & RAG)")

# --- VARIABLES DE LÍMITES DE NEGOCIO ---
LIMIT_BASICO_ANALYSIS_DAILY = int(os.getenv("LIMIT_BASICO_ANALYSIS_DAILY", 3))
LIMIT_AVANZADO_ANALYSIS_DAILY = int(os.getenv("LIMIT_AVANZADO_ANALYSIS_DAILY", 15))
LIMIT_PREMIUM_ANALYSIS_DAILY = int(os.getenv("LIMIT_PREMIUM_ANALYSIS_DAILY", 25))

LIMIT_BASICO_DOCS = int(os.getenv("LIMIT_BASICO_DOCS", 1))
LIMIT_AVANZADO_DOCS = int(os.getenv("LIMIT_AVANZADO_DOCS", 3))
LIMIT_PREMIUM_DOCS = int(os.getenv("LIMIT_PREMIUM_DOCS", 5))

LIMIT_SIZE_MB_BASICO = int(os.getenv("LIMIT_SIZE_MB_BASICO", 10))
LIMIT_SIZE_MB_AVANZADO = int(os.getenv("LIMIT_SIZE_MB_AVANZADO", 50))
LIMIT_SIZE_MB_PREMIUM = int(os.getenv("LIMIT_SIZE_MB_PREMIUM", 50))
LIMIT_SIZE_MB_VIP = int(os.getenv("LIMIT_SIZE_MB_VIP", 50))

# --- CORS ---
raw_origins = os.getenv("ALLOWED_ORIGINS", '["[https://pida-ai.com](https://pida-ai.com)"]')
try:
    origins = json.loads(raw_origins)
except:
    origins = ["[https://pida-ai.com](https://pida-ai.com)"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_origin_regex=r"https://pida-ai-v20--.*\.web\.app$|https://.*\.app\.github\.dev$",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

# --- MAPAS DE LÍMITES DE NEGOCIO ---
ANALYSIS_LIMITS = {
    "basico": LIMIT_BASICO_ANALYSIS_DAILY, "avanzado": LIMIT_AVANZADO_ANALYSIS_DAILY,
    "premium": LIMIT_PREMIUM_ANALYSIS_DAILY, "vip": -1 
}
DOCS_LIMITS = {
    "basico": LIMIT_BASICO_DOCS, "avanzado": LIMIT_AVANZADO_DOCS,
    "premium": LIMIT_PREMIUM_DOCS, "vip": 100 
}
PLAN_SIZE_LIMITS = {
    "basico": LIMIT_SIZE_MB_BASICO, "avanzado": LIMIT_SIZE_MB_AVANZADO,
    "premium": LIMIT_SIZE_MB_PREMIUM, "vip": LIMIT_SIZE_MB_VIP
}

# --- FUNCIONES DE UTILIDAD Y CONTROL ---
def get_date_utc_minus_6() -> str:
    utc_now = datetime.now(timezone.utc)
    cst_now = utc_now - timedelta(hours=6)
    return cst_now.strftime('%Y-%m-%d')

async def get_user_plan_unified(current_user: Dict[str, Any]) -> str:
    user_id = current_user.get('uid')
    user_email = current_user.get('email', '').strip().lower()
    try:
        raw_domains = os.getenv("ADMIN_DOMAINS", '[]')
        raw_emails = os.getenv("ADMIN_EMAILS", '[]')
        admin_domains = [str(d).strip().lower() for d in json.loads(raw_domains)]
        admin_emails = [str(e).strip().lower() for e in json.loads(raw_emails)]
    except:
        admin_domains, admin_emails = [], []

    email_domain = user_email.split("@")[-1] if "@" in user_email else ""
    if (email_domain in admin_domains) or (user_email in admin_emails): return 'vip'

    try:
        cust_doc = await db.collection('customers').document(user_id).get()
        if cust_doc.exists:
            data = cust_doc.to_dict()
            if data.get('status') in ['active', 'trialing']:
                return data.get('plan', 'basico').lower()
    except Exception as e:
        print(f"Error consultando plan en DB: {e}")
    return 'none'

async def consume_analysis_credit(user_id: str, plan_key: str):
    limit_daily = ANALYSIS_LIMITS.get(plan_key, 0)
    if limit_daily == -1: return 
    today = get_date_utc_minus_6()
    stats_ref = db.collection('users').document(user_id).collection('usage_stats').document(today)
    
    @firestore.async_transactional
    async def check_and_increment(transaction, ref):
        snapshot = await ref.get(transaction=transaction)
        data = snapshot.to_dict() if snapshot.exists else {}
        current_count = data.get('analysis_count', 0)
        if current_count >= limit_daily:
            raise HTTPException(status_code=429, detail=f"Límite diario alcanzado para el plan {plan_key}")
        transaction.set(ref, {'analysis_count': current_count + 1, 'last_updated': firestore.SERVER_TIMESTAMP}, merge=True)
    transaction = db.transaction()
    await check_and_increment(transaction, stats_ref)

async def refund_analysis_credit(user_id: str):
    today = get_date_utc_minus_6()
    stats_ref = db.collection('users').document(user_id).collection('usage_stats').document(today)
    @firestore.async_transactional
    async def check_and_decrement(transaction, ref):
        snapshot = await ref.get(transaction=transaction)
        if snapshot.exists:
            current_count = (snapshot.to_dict() or {}).get('analysis_count', 0)
            if current_count > 0:
                transaction.update(ref, {'analysis_count': current_count - 1, 'last_updated': firestore.SERVER_TIMESTAMP})
    try:
        transaction = db.transaction()
        await check_and_decrement(transaction, stats_ref)
    except Exception as e: print(f"Error en reembolso: {e}")

def generate_filename(instructions: str, extension: str) -> str:
    safe_title = re.sub(r'[^a-zA-Z0-9áéíóúÁÉÍÓÚñÑ ]', '', instructions[:40]).strip().replace(' ', '_')
    return f"{safe_title or 'Analisis_PIDA'}_{datetime.now().strftime('%Y-%m-%d-%H-%M-%S')}.{extension}"

def sanitize_text_for_pdf(text: str) -> str:
    if not text: return ""
    replacements = {"•": "-", "—": "-", "–": "-", "“": '"', "”": '"', "‘": "'", "’": "'", "…": "...", "\u2013": "-", "\u2014": "-", "\u2022": "-", "\uF0B7": "-"}
    for char, rep in replacements.items(): text = text.replace(char, rep)
    return text.encode('latin1', 'replace').decode('latin-1')

def get_multi_cell_height(pdf, w, text):
    lines = 0
    for paragraph in str(text).split('\n'):
        words = paragraph.split(' ')
        if not words or (len(words) == 1 and words[0] == ''):
            lines += 1; continue
        current_line = ""
        for word in words:
            if pdf.get_string_width(current_line + word + " ") > w and current_line:
                lines += 1; current_line = word + " "
            else: current_line += word + " "
        lines += 1
    return lines * 6

def render_table_row_fpdf(pdf, cells):
    col_width = (pdf.w - pdf.l_margin - pdf.r_margin) / max(len(cells), 1)
    processed_cells = [c.replace('<br>', '\n').replace('<br/>', '\n').replace('**', '').strip() for c in cells]
    max_height = max([get_multi_cell_height(pdf, col_width, cell) for cell in processed_cells] + [0])
    row_height = max_height if max_height > 0 else 6
    
    if pdf.get_y() + row_height > pdf.page_break_trigger: pdf.add_page()
    start_x, start_y = pdf.get_x(), pdf.get_y()
    
    for cell in processed_cells:
        x_before, y_before = pdf.get_x(), pdf.get_y()
        pdf.rect(x_before, y_before, col_width, row_height)
        pdf.multi_cell(col_width, 6, cell, border=0, align='L')
        pdf.set_xy(x_before + col_width, start_y)
    pdf.set_xy(pdf.l_margin, start_y + row_height)

def write_markdown_to_pdf(pdf, text):
    import re
    pdf.set_font("Arial", "", 11)
    
    try:
        effective_page_width = pdf.epw
    except AttributeError:
        effective_page_width = pdf.w - pdf.l_margin - pdf.r_margin
        
    text = text.replace('$', '').replace('^{a}', 'a.').replace('^{o}', 'o.')
    
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # --- 1. PROCESAMIENTO DE TABLAS MARKDOWN ---
        if line.startswith('|') and line.endswith('|'):
            table_lines = []
            # Agrupar todas las líneas de la tabla antes de dibujar
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
                
            for r_idx, t_line in enumerate(table_lines):
                cols = [c.strip() for c in t_line.split('|')[1:-1]]
                
                # Ignorar filas separadoras (ej. |---|---|)
                if all(re.match(r'^:?-+:?$', c) for c in cols):
                    continue
                if not cols:
                    continue
                    
                col_width = effective_page_width / len(cols)
                
                # Función interna para calcular la altura dinámica de la celda
                def get_cell_height(w, txt, is_bold):
                    pdf.set_font("Arial", "B" if is_bold else "", 10)
                    try: margin = pdf.c_margin
                    except: margin = 1
                    usable_w = w - (2 * margin) # Restar márgenes internos de FPDF
                    
                    lines_count = 0
                    for p in str(txt).split('\n'):
                        words = p.split(' ')
                        if not words or (len(words) == 1 and words[0] == ''):
                            lines_count += 1; continue
                        curr_line = ""
                        for word in words:
                            if pdf.get_string_width(curr_line + word + " ") > usable_w and curr_line:
                                lines_count += 1; curr_line = word + " "
                            else: 
                                curr_line += word + " "
                        lines_count += 1
                    return lines_count * 6
                
                is_header = (r_idx == 0)
                # Calcular la altura de la fila basándose en la celda con más texto
                max_height = max([get_cell_height(col_width, c.replace('**', '').replace('<br>', '\n').replace('<br/>', '\n'), is_bold=(is_header or '**' in c)) for c in cols] + [6])
                
                # Prevenir salto de página a la mitad de una fila
                try: pb_trigger = pdf.page_break_trigger
                except: pb_trigger = pdf.h - pdf.b_margin
                if pdf.get_y() + max_height > pb_trigger:
                    pdf.add_page()
                    
                x_start = pdf.get_x()
                y_start = pdf.get_y()
                
                # Dibujar las celdas
                for c_idx, col in enumerate(cols):
                    col_clean = col.replace('**', '')
                    col_clean = re.sub(r'<br\s*/?>', '\n', col_clean, flags=re.IGNORECASE)
                    
                    # Dibujar fondo y contorno FIRST
                    if is_header:
                        pdf.set_fill_color(241, 245, 249) # Azul muy claro tipo Tailwind (#f1f5f9)
                        pdf.rect(x_start + (c_idx * col_width), y_start, col_width, max_height, 'DF')
                    else:
                        pdf.rect(x_start + (c_idx * col_width), y_start, col_width, max_height)
                    
                    # Imprimir el texto encima
                    pdf.set_xy(x_start + (c_idx * col_width), y_start)
                    pdf.set_font("Arial", "B" if "**" in col or is_header else "", 10)
                    
                    if is_header:
                        pdf.set_text_color(29, 53, 87) # Azul marino
                    else:
                        pdf.set_text_color(0, 0, 0)
                        
                    pdf.multi_cell(col_width, 6, col_clean, border=0, align='L')
                
                # Acomodar el cursor debajo de la fila recién dibujada
                pdf.set_xy(x_start, y_start + max_height)
            
            # Restaurar colores y fuentes al terminar la tabla
            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(5)
            continue

        # --- 2. ESPACIOS VACÍOS ---
        if not line:
            pdf.ln(5)
            i += 1
            continue

        # --- 3. ENCABEZADOS PRINCIPALES (##) ---
        if line.startswith('## '):
            pdf.ln(3)
            pdf.set_font("Arial", "B", 13)
            pdf.set_text_color(29, 53, 87)
            clean_line = re.sub(r'<br\s*/?>', ' ', line.replace('## ', ''), flags=re.IGNORECASE)
            pdf.multi_cell(0, 8, clean_line)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)
            i += 1
            continue
            
        # --- 4. SUBTÍTULOS (###) ---
        if line.startswith('### '):
            pdf.ln(2)
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(40, 70, 100)
            clean_line = re.sub(r'<br\s*/?>', ' ', line.replace('### ', ''), flags=re.IGNORECASE)
            pdf.multi_cell(0, 7, clean_line)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)
            i += 1
            continue

        # --- 5. LISTAS Y PÁRRAFOS REGULARES ---
        if line.startswith('* ') or line.startswith('- '):
            pdf.set_x(15)
            line = "- " + line[2:]
        else:
            pdf.set_x(10)

        line = re.sub(r'<br\s*/?>', '', line, flags=re.IGNORECASE)

        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                pdf.set_font("Arial", "B", 11)
                pdf.write(6, part.strip('*'))
                pdf.set_font("Arial", "", 11)
            else:
                pdf.write(6, part)
        pdf.ln(6)
        
        i += 1
def parse_and_add_markdown_to_docx(document, markdown_text):
    lines = markdown_text.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('|') and line.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip()); i += 1
            if len(table_lines) > 0:
                headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                table = document.add_table(rows=1, cols=len(headers))
                try: table.style = 'Table Grid'
                except: pass
                for col_idx, header in enumerate(headers):
                    if col_idx < len(table.rows[0].cells):
                        cell_text = header.replace('**', '').replace('<br>', '\n').replace('<br/>', '\n')
                        # En lugar de asignar el texto directo, lo hacemos con un "run" para poner negrita
                        p = table.rows[0].cells[col_idx].paragraphs[0]
                        p.text = "" 
                        run = p.add_run(cell_text)
                        run.bold = True
                
                start_idx = 2 if len(table_lines) > 1 and all(re.match(r'^:?-+:?$', c.strip()) for c in table_lines[1].split('|')[1:-1]) else 1
                for r_idx in range(start_idx, len(table_lines)):
                    row_cells_text = [c.strip() for c in table_lines[r_idx].split('|')[1:-1]]
                    row = table.add_row()
                    for col_idx, cell_text in enumerate(row_cells_text):
                        if col_idx < len(row.cells):
                            row.cells[col_idx].text = cell_text.replace('**', '').replace('<br>', '\n').replace('<br/>', '\n')
            continue
            
        if line.startswith('### '): document.add_heading(line.lstrip('### '), level=3)
        elif line.startswith('## '): document.add_heading(line.lstrip('## '), level=2)
        elif line.startswith('# '): document.add_heading(line.lstrip('# '), level=1)
        elif not line: document.add_paragraph('')
        else:
            p = document.add_paragraph()
            for part in re.split(r'(\*\*.*?\*\*)', line):
                if part.startswith('**') and part.endswith('**'): p.add_run(part.strip('*')).bold = True
                else: p.add_run(part)
        i += 1

class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 14); self.set_text_color(29, 53, 87)
        self.cell(0, 10, "PIDA-AI: Resumen de Consulta", 0, 1, "L")
        self.set_font("Arial", "", 9); self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Generado: {datetime.now().strftime('%d/%m/%Y, %H:%M:%S')}", 0, 1, "L"); self.ln(5)
    def footer(self):
        self.set_y(-15); self.set_font("Arial", "", 8); self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Pagina {self.page_no()}/{{nb}}", 0, 0, "C")

def read_docx_sync(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
        return "\n".join([p.text for p in doc.paragraphs])
    except: return ""

def download_and_parse_docx(gs_uri: str) -> str:
    try:
        blob_path = "/".join(gs_uri.replace("gs://", "").split("/")[1:])
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        content = bucket.blob(blob_path).download_as_bytes()
        return read_docx_sync(content)
    except Exception as e:
        print(f"Error procesando DOCX: {e}")
        return ""

def create_docx_sync(analysis_text: str, instructions: str) -> tuple[bytes, str, str]:
    stream = io.BytesIO()
    doc = Document()
    doc.add_heading("PIDA-AI: Resumen", 0)
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_heading("Instrucciones", 2)
    doc.add_paragraph(instructions)
    doc.add_heading("Analisis", 2)
    
    parse_and_add_markdown_to_docx(doc, analysis_text)

    doc.save(stream); stream.seek(0)
    return stream.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", generate_filename(instructions, "docx")

def create_pdf_sync(analysis_text: str, instructions: str) -> tuple[bytes, str, str]:
    safe_inst, safe_ana = sanitize_text_for_pdf(instructions), sanitize_text_for_pdf(analysis_text)
    pdf = PDF(); pdf.alias_nb_pages(); pdf.add_page()
    pdf.set_font("Arial", "B", 12); pdf.cell(0, 10, "Instrucciones", 0, 1); pdf.set_font("Arial", "", 11); pdf.multi_cell(0, 6, safe_inst); pdf.ln(5)
    pdf.set_font("Arial", "B", 12); pdf.cell(0, 10, "Analisis", 0, 1)
    
    if not safe_ana.strip():
        pdf.set_font("Arial", "I", 11); pdf.multi_cell(0, 6, "[Sin contenido]")
    else:
        write_markdown_to_pdf(pdf, safe_ana)

    try:
        pdf_string = pdf.output(dest='S')
        pdf_bytes = pdf_string.encode('latin-1', 'replace') if isinstance(pdf_string, str) else pdf_string
        stream = io.BytesIO(pdf_bytes)
        return stream.read(), "application/pdf", generate_filename(instructions, "pdf")
    except Exception as e:
        err = FPDF(); err.add_page(); err.multi_cell(0, 10, f"Error: {str(e)}")
        return err.output(dest='S').encode('latin-1'), "application/pdf", "Error.pdf"


def format_json_visualizations_for_export(text: str) -> str:
    """
    Busca los bloques de JSON nativos en el texto de análisis y los 
    convierte en Tablas Markdown para que las funciones de PDF y DOCX
    puedan dibujarlos correctamente.
    """
    if not text:
        return ""
        
    # --- CONVERTIR TIMELINE EN TABLA ---
    def replace_timeline(match):
        raw_json = match.group(1).strip()
        raw_json = re.sub(r'`{3}(?:json-timeline|json)?', '', raw_json, flags=re.IGNORECASE).replace('`'*3, '').strip()
        try:
            data = json.loads(raw_json)
            if not data: return ""
            table = "\n### Línea de Tiempo\n\n| Fecha | Fase | Descripción |\n|---|---|---|\n"
            for item in data:
                date = str(item.get('date', '')).replace('\n', ' ').strip()
                phase = str(item.get('phase', '')).replace('\n', ' ').strip()
                desc = str(item.get('description', '')).replace('\n', ' ').strip()
                table += f"| **{date}** | {phase} | {desc} |\n"
            return table + "\n"
        except Exception as e:
            print(f"Error exportando Timeline: {e}")
            return match.group(0)

    text = re.sub(r'\[TIMELINE_START\](.*?)\[TIMELINE_END\]', replace_timeline, text, flags=re.DOTALL | re.IGNORECASE)

    # --- CONVERTIR FLOW EN TABLA ---
    def replace_flow(match):
        raw_json = match.group(1).strip()
        raw_json = re.sub(r'`{3}(?:json-flow|json)?', '', raw_json, flags=re.IGNORECASE).replace('`'*3, '').strip()
        try:
            data = json.loads(raw_json)
            if not data: return ""
            table = "\n### Diagrama de Flujo / Proceso\n\n| Paso | Requisito | Acción |\n|---|---|---|\n"
            for i, item in enumerate(data):
                step = str(item.get('step', '')).replace('\n', ' ').strip()
                req = str(item.get('requirement', '')).replace('\n', ' ').strip()
                action = str(item.get('action', '')).replace('\n', ' ').strip()
                table += f"| **{i+1}. {step}** | {req} | {action} |\n"
            return table + "\n"
        except Exception as e:
            print(f"Error exportando Flow: {e}")
            return match.group(0)

    text = re.sub(r'\[FLOW_START\](.*?)\[FLOW_END\]', replace_flow, text, flags=re.DOTALL | re.IGNORECASE)

    return text


# --- CORE GENAI STREAMING GENERATOR ---
async def stream_analysis_generator(genai_client, model_name, contents, gen_config, current_user, instructions, original_filenames, files_info, analysis_id, db_history):
    full_text = ""
    try:
        # Llamada Asíncrona con el SDK construida con soporte de chat/historial
        responses = await genai_client.aio.models.generate_content_stream(
            model=model_name,
            contents=contents,
            config=gen_config
        )
        
        async for chunk in responses:
            if chunk.text:
                full_text += chunk.text
                yield f"data: {json.dumps({'text': chunk.text})}\n\n"
        
        user_id = current_user.get("uid")
        
        # Registrar respuesta del modelo
        db_history.append({"role": "model", "content": full_text})
        final_id = analysis_id
        
        if final_id:
            doc_ref = db.collection("analysis_history").document(final_id)
            doc = await doc_ref.get()
            if doc.exists:
                await doc_ref.update({
                    "analysis": json.dumps(db_history), 
                    "timestamp": SERVER_TIMESTAMP
                })
            else: final_id = "" 
                
        if not final_id:
            title_source = db_history[0].get("content", instructions) if db_history else instructions
            title = (title_source[:40] + '...') if len(title_source) > 40 else title_source
            doc_ref = db.collection("analysis_history").document()
            await doc_ref.set({
                "userId": user_id, 
                "title": title, 
                "instructions": title_source,
                "analysis": json.dumps(db_history), 
                "timestamp": SERVER_TIMESTAMP, 
                "original_filenames": original_filenames,
                "files_data": files_info # <- Blindaje: Se almacena metadata de archivos en Backend
            })
            final_id = doc_ref.id
        
        yield f"data: {json.dumps({'done': True, 'analysis_id': final_id})}\n\n"
        
    except APIError as e:
        print(f"Error GenAI API: {e}")
        yield f"data: {json.dumps({'error': f'Error del modelo: {e.message}'})}\n\n"
    except Exception as e:
        print(f"Error stream: {e}")
        yield f"data: {json.dumps({'error': str(e)})}\n\n"

# --- ENDPOINTS ---

@app.post("/compress-and-upload")
async def compress_and_upload(file: UploadFile = File(...), current_user: Dict[str, Any] = Depends(get_current_user)):
    user_id = current_user['uid']
    plan = await get_user_plan_unified(current_user)
    if plan == 'none': raise HTTPException(403, "Plan inactivo.")
    max_size_mb = PLAN_SIZE_LIMITS.get(plan, 10)
    file_bytes = await file.read()
    original_size = len(file_bytes) / (1024 * 1024)

    if original_size > max_size_mb: raise HTTPException(400, f"EXCEDE_TAMANO: El archivo pesa {original_size:.2f} MB. Max permitido: {max_size_mb} MB.")

    try:
        def compress_pdf(data: bytes) -> bytes:
            if file.content_type != "application/pdf": return data
            doc = fitz.open(stream=data, filetype="pdf")
            return doc.tobytes(garbage=4, deflate=True)

        compressed_bytes = await asyncio.to_thread(compress_pdf, file_bytes)
        new_size = len(compressed_bytes) / (1024 * 1024)
        bucket = storage_client.bucket(GCS_BUCKET_NAME)
        safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '', file.filename.replace(' ', '_'))
        blob_path = f"uploads/{user_id}/{int(datetime.now().timestamp())}_opt_{safe_name}"
        blob = bucket.blob(blob_path)
        await asyncio.to_thread(blob.upload_from_string, compressed_bytes, content_type=file.content_type)

        return {
            "filename": file.filename, "gs_uri": f"gs://{GCS_BUCKET_NAME}/{blob_path}",
            "mime_type": file.content_type, "original_size_mb": original_size, "new_size_mb": new_size
        }
    except Exception as e: raise HTTPException(500, f"Error optimizando: {str(e)}")

@app.post("/generate-upload-urls")
async def generate_upload_urls(payload: Dict[str, Any] = Body(...), current_user: Dict[str, Any] = Depends(get_current_user)):
    user_id = current_user['uid']
    plan = await get_user_plan_unified(current_user)
    if plan == 'none': raise HTTPException(403, "Plan inactivo.")

    files_req = payload.get("files", [])
    max_docs = DOCS_LIMITS.get(plan, 0)
    if max_docs != -1 and len(files_req) > max_docs: raise HTTPException(403, f"Tu plan permite {max_docs} documento(s) a la vez.")

    urls_response = []
    bucket = storage_client.bucket(GCS_BUCKET_NAME)
    for f in files_req:
        safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '', f.get("name", "doc").replace(' ', '_'))
        blob_path = f"uploads/{user_id}/{int(datetime.now().timestamp())}_{safe_name}"
        blob = bucket.blob(blob_path)
        try:
            signed_url = blob.generate_signed_url(version="v4", expiration=timedelta(minutes=15), method="PUT", content_type=f.get("type", "application/pdf"))
            urls_response.append({"filename": f.get("name"), "upload_url": signed_url, "gs_uri": f"gs://{GCS_BUCKET_NAME}/{blob_path}", "mime_type": f.get("type", "application/pdf")})
        except Exception as e: raise HTTPException(500, f"Error generando URL segura: {e}")
    return {"urls": urls_response}

@app.post("/analyze")
async def analyze_documents(
    files_data: str = Form("[]"), instructions: str = Form(...), analysis_id: str = Form(""),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    if not genai_client: raise HTTPException(500, "El cliente de IA no está inicializado.")
    
    user_id = current_user['uid']
    plan = await get_user_plan_unified(current_user)
    if plan == 'none': raise HTTPException(403, "No tienes un plan activo.")

    # 1. Recuperar Estado (Historial y Archivos) desde Firestore
    db_history = []
    original_filenames = []
    files_info = []
    is_follow_up = False

    if analysis_id:
        doc_ref = db.collection("analysis_history").document(analysis_id)
        doc = await doc_ref.get()
        if doc.exists:
            doc_data = doc.to_dict()
            if doc_data.get("userId") != user_id:
                raise HTTPException(403, "No tienes permiso para acceder a este historial.")
            
            db_history = json.loads(doc_data.get("analysis", "[]"))
            files_info = doc_data.get("files_data", [])
            original_filenames = doc_data.get("original_filenames", [])
            is_follow_up = True
        else:
            analysis_id = "" 

    if not is_follow_up:
        try: files_info = json.loads(files_data)
        except: raise HTTPException(400, "Formato de metadatos inválido.")
        if DOCS_LIMITS.get(plan, 0) != -1 and len(files_info) > DOCS_LIMITS.get(plan, 0): 
            raise HTTPException(403, f"Plan excede límite de documentos.")
            
    await consume_analysis_credit(user_id, plan)

    max_size_mb = PLAN_SIZE_LIMITS.get(plan, 10)
    model_parts = []
    bucket = storage_client.bucket(GCS_BUCKET_NAME)

    # 2. Reconstruir los documentos (model_parts)
    for f_info in files_info:
        gs_uri = f_info.get("gs_uri")
        mime_type = f_info.get("mime_type", "application/pdf")
        original_filename = f_info.get("filename", "documento")
        
        try:
            blob = bucket.get_blob("/".join(gs_uri.replace("gs://", "").split("/")[1:]))
            if blob and (blob.size / (1024 * 1024)) > max_size_mb:
                blob.delete()
                raise HTTPException(400, f"EXCEDE_TAMANO: '{original_filename}' excede los {max_size_mb} MB.")
        except HTTPException as he: raise he
        except Exception as e: print(f"Error validando archivo en GCS: {e}")

        if not is_follow_up:
            original_filenames.append(original_filename)
        
        if mime_type == "application/pdf":
            model_parts.append(types.Part.from_uri(file_uri=gs_uri, mime_type="application/pdf"))
        else:
            text = await asyncio.to_thread(download_and_parse_docx, gs_uri)
            model_parts.append(types.Part.from_text(text=f"--- DOC: {original_filename} ---\n{text}\n------\n"))

    # =====================================================================
    # NUEVO: LÓGICA DE INTEGRACIÓN CON RAG INTERNO (TOLERANTE A FALLOS)
    # =====================================================================
    rag_context_text = ""
    if not is_follow_up: # Idealmente solo lo consultamos en la pregunta inicial para no saturar
        try:
            rag_url = os.getenv("RAG_API_URL")
            if rag_url:
                print("Consultando RAG interno para enriquecer contexto del analizador...")
                timeout_config = httpx.Timeout(15.0, connect=5.0)
                
                async with httpx.AsyncClient(timeout=timeout_config) as http_client:
                    resp = await http_client.post(rag_url, json={"query": instructions})
                    if resp.status_code == 200:
                        rag_data = resp.json()
                        if rag_data and "results" in rag_data and rag_data["results"]:
                            rag_context_text = "\n\n### Contexto de Documentos Internos (RAG):\n"
                            
                            for doc in rag_data.get("results", []):
                                title = doc.get("title")
                                author = doc.get("author")
                                source_filename = doc.get("source")
                                content = doc.get("content", "").replace("\n", " ").strip()
                                
                                display_title = title or source_filename or "Documento Interno"
                                citation_line = f"Título: {display_title}"
                                if author and author.strip() and author != "Autor Desconocido":
                                    citation_line += f" | Autor: {author}"
                                else:
                                    citation_line += f" | Autor: Institucional/No especificado"
                                
                                rag_context_text += f"{citation_line}\n**Texto:**\n> {content}\n\n"
                            print("Contexto RAG obtenido exitosamente.")
        except Exception as e:
            print(f"Error consultando RAG (silenciado para no afectar análisis): {e}")
    # =====================================================================

    # 3. Construir lista Nativa de Contenidos (Chat History)
    contents = []
    
    if not is_follow_up:
        db_history.append({"role": "user", "content": instructions})
        
        # Usamos el prompt centralizado de primer turno
        first_turn_prompt = FIRST_TURN_PROMPT_TEMPLATE.format(instructions=instructions)
        
        # 👇 INYECTAMOS EL RAG AQUÍ
        if rag_context_text:
            first_turn_prompt += f"\n\nInstrucción adicional: Toma en cuenta el siguiente contexto de jurisprudencia y documentos internos de PIDA para complementar tu análisis si es relevante a la pregunta.\n{rag_context_text}"
            
        user_parts = model_parts + [types.Part.from_text(text=first_turn_prompt)]
        contents.append(types.Content(role="user", parts=user_parts))
    else:
        for i, msg in enumerate(db_history):
            role = msg.get("role", "user")
            text_part = types.Part.from_text(text=msg.get("content", ""))
            
            # Anclar los documentos únicamente en el primer mensaje
            if i == 0 and role == "user":
                contents.append(types.Content(role=role, parts=model_parts + [text_part]))
            else:
                contents.append(types.Content(role=role, parts=[text_part]))
                
        # IMPORTANTE: Usamos la plantilla centralizada para que SIEMPRE lleve las reglas del JSON
        # Y NO duplicamos db_history.append aquí (ya se hace en el stream_generator)
        follow_up_prompt = FOLLOW_UP_PROMPT_TEMPLATE.format(instructions=instructions)
        contents.append(types.Content(role="user", parts=[types.Part.from_text(text=follow_up_prompt)]))
        
        # FIX: Guardar la nueva pregunta del usuario en la base de datos para follow-ups
        db_history.append({"role": "user", "content": instructions})

    gen_config = types.GenerateContentConfig(
        system_instruction=ANALYZER_SYSTEM_PROMPT,
        temperature=float(os.getenv("GEMINI_TEMP", "0.4")),
        top_p=float(os.getenv("GEMINI_TOP_P", "0.95")),
        max_output_tokens=65535,
        safety_settings=[
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_ONLY_HIGH"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_ONLY_HIGH"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_ONLY_HIGH"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_ONLY_HIGH"),
        ]
    )

    async def counted_stream_generator():
        has_error = False; tokens_sent = False
        try:
            # 👇 AQUÍ ESTÁ EL LATIDO: Se envía inmediatamente para mantener viva la conexión QUIC/TCP
            yield f"data: {json.dumps({'status': 'Analizando contenido extenso. Esto puede tomar un momento...'})}\n\n"
            
            async for chunk in stream_analysis_generator(
                genai_client, GEMINI_MODEL_NAME, contents, gen_config, current_user, instructions, original_filenames, files_info, analysis_id, db_history
            ):
                if '"error":' in chunk: has_error = True
                if '"text":' in chunk and not has_error: tokens_sent = True
                yield chunk
        finally:
            if has_error or not tokens_sent: asyncio.create_task(refund_analysis_credit(user_id))

    # =====================================================================
    # NUEVO: REGISTRO DE ESTADÍSTICA MENSUAL (AHORRO DE LECTURAS)
    # =====================================================================
    try:
        current_month = datetime.now().strftime("%Y-%m")
        stats_ref = db.collection('monthly_stats').document(current_month)
        await stats_ref.set({
            "analisis": firestore.Increment(1)
        }, merge=True)
    except Exception as stats_e:
        print(f"Error guardando estadística mensual de análisis: {stats_e}")
    # =====================================================================

    # 👇 HEADERS ANTIBÚFER: Para que el latido y el texto salgan al instante
    headers = { 
        "Content-Type": "text/event-stream", 
        "Cache-Control": "no-cache", 
        "Connection": "keep-alive", 
        "X-Accel-Buffering": "no" 
    }
    
    return StreamingResponse(counted_stream_generator(), headers=headers)

@app.post("/download-analysis")
async def download_analysis(
    analysis_text: str = Form(""), instructions: str = Form("Exportación de Análisis PIDA"),
    history_json: str = Form(""), file_format: str = Form("docx"), analysis_id: Optional[str] = Form(None),
    current_user: Dict[str, Any] = Depends(get_current_user)
):
    if await get_user_plan_unified(current_user) == 'none': raise HTTPException(403, "Sin acceso")
    if history_json:
        try:
            messages = json.loads(history_json)
            chat_lines = []
            for idx, msg in enumerate(messages):
                role, content = msg.get("role"), msg.get("content", "")
                if idx == 0 and role == "user": instructions = content
                else: chat_lines.append(f"**{'Instrucción' if role == 'user' else 'Análisis PIDA'}:**\n{content}")
            analysis_text = "\n\n".join(chat_lines)
        except Exception: pass

    # Transformar los bloques JSON a tablas Markdown antes de convertirlos
    analysis_text = format_json_visualizations_for_export(analysis_text)

    analysis_text = analysis_text[:500000] + "\n\n[Texto truncado]" if len(analysis_text) > 500000 else analysis_text
    instructions = instructions[:5000] + "..." if len(instructions) > 5000 else instructions

    try:
        content, mime, fname = await asyncio.to_thread(create_docx_sync if file_format.lower() == "docx" else create_pdf_sync, analysis_text, instructions)
        return Response(content=content, media_type=mime, headers={"Content-Disposition": f"attachment; filename={fname}"})
    except Exception as e: raise HTTPException(500, f"Error descarga: {e}")

@app.get("/analysis-history/")
async def get_analysis_history(current_user: Dict[str, Any] = Depends(get_current_user)):
    user_id = current_user['uid']
    if await get_user_plan_unified(current_user) == 'none': raise HTTPException(403, "Requiere plan activo")
    history = []
    async for d in db.collection("analysis_history").where("userId", "==", user_id).order_by("timestamp", direction=Query.DESCENDING).stream():
        history.append({"id": d.id, "title": d.get("title"), "timestamp": d.get("timestamp"), "userId": user_id})
    return history

@app.get("/analysis-history/{analysis_id}")
async def get_analysis_detail(analysis_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    if await get_user_plan_unified(current_user) == 'none': raise HTTPException(403)
    doc = await db.collection("analysis_history").document(analysis_id).get()
    if not doc.exists or doc.to_dict().get("userId") != current_user['uid']: raise HTTPException(404)
    return doc.to_dict()

@app.delete("/analysis-history/{analysis_id}")
async def delete_analysis(analysis_id: str, current_user: Dict[str, Any] = Depends(get_current_user)):
    doc_ref = db.collection("analysis_history").document(analysis_id)
    doc = await doc_ref.get()
    if not doc.exists or doc.to_dict().get("userId") != current_user['uid']: raise HTTPException(404)
    await doc_ref.delete()
    return {"status": "ok"}

@app.get("/")
def read_root(): return {"status": "ok", "msg": "API Analizador GenAI v3.0"}
